# Speech to text imports
from vosk import Model, KaldiRecognizer, SetLogLevel
from pydub import AudioSegment
# PDF imports
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph
from reportlab.platypus import SimpleDocTemplate
# DocX imports
from docx import Document
# Grammar fixer import
import language_tool_python
# Class import
import Summarizer as Sum
# Base imports
import sys
import os
import wave
import json
import threading
import time
# GUI imports
import tkinter as tk
import tkinter.ttk as ttk
from tkinter import filedialog


filepath = ""
top = tk.Tk()
top.title("LectureAid")
top.geometry("977x700")
top.configure(background='#242320')
frame = ttk.Frame()
frame.place(relx=0.5, rely=0.5, relwidth=0.80,relheight=0.4, anchor=tk.CENTER)
frame_style = ttk.Style()
frame_style.configure('TFrame', background='#242320')


def exit_function():
    top.destroy()
    sys.exit(0)


top.protocol('WM_DELETE_WINDOW', exit_function)

fileLabel = tk.Label(frame, text="Selected file: None", background='#242320', foreground='white', font=('Arial', 12), pady=5)
fileLabel.pack(side=tk.TOP)

var = tk.IntVar()

R1 = tk.Radiobutton(frame, text="Full Transcription (DEFAULT)", variable=var, value=1, background='#3b3a37', foreground='white',  selectcolor='green', pady=5)
R2 = tk.Radiobutton(frame, text="Summarized Transcription", variable=var, value=2, background='#3b3a37', foreground='white',  selectcolor='green', pady=5)
R3 = tk.Radiobutton(frame, text="Full+Summarized Transcription", variable=var, value=3, background='#3b3a37', foreground='white', selectcolor='green', pady=5)

def open_file():
    global filepath
    filepath = filedialog.askopenfilename(initialdir="/", title="Select a File", filetypes=(("wav files", "*.wav"), ("mp3 files", "*.mp3"), ("mp4 files", "*.mp4"),))
    fileLabel['text'] = "Selected file: " + filepath
    startButton.pack(side=tk.TOP)
    R1.pack(side=tk.TOP)
    R2.pack(side=tk.TOP)
    R3.pack(side=tk.TOP)


fileButton = tk.Button(frame, text="Select a File", command=open_file, background='#3b3a37', foreground='white', width='35', pady=5)
fileButton.pack(side=tk.TOP)


def SpeechToText():
    SetLogLevel(0)
    if not os.path.exists("model"):
        print(
            "Please download the model from https://alphacephei.com/vosk/models and unpack as 'model' in the current folder.")
        exit(1)

    startButton.pack_forget()
    fileButton.pack_forget()
    R1.pack_forget()
    R2.pack_forget()
    R3.pack_forget()
    global filepath
    print(filepath)

    total = AudioSegment.from_file(filepath)
    total = total.set_frame_rate(16000)
    total = total.set_channels(1)
    total = total.set_sample_width(2)
    total.export("chunks.wav", format="wav")

    totalTimeLabel = tk.Label(frame, text="Total time of file: " + time.strftime("%H:%M:%S", time.gmtime(len(total)/1000)),background='#242320', foreground='white', font=('Arial', 12), pady=5)
    totalTimeLabel.pack(side=tk.TOP)
    percentLabel = tk.Label(frame, text="% Converted: 0", background='#242320', foreground='white', font=('Arial', 12), pady=5)
    percentLabel.pack(side=tk.TOP)

    wf = wave.open("chunks.wav", "rb")
    if wf.getnchannels() != 1 or wf.getsampwidth() != 2 or wf.getcomptype() != "NONE":
        print("Audio file must be WAV format mono PCM.")
        exit(1)

    model = Model("model")
    rec = KaldiRecognizer(model, wf.getframerate())

    allText = ""
    toSummary = ""
    i = 0
    while True:
        data = wf.readframes(4000)
        if len(data) == 0:
            break
        if rec.AcceptWaveform(data):
            i += 1
            print(rec.Result())
            result = json.loads(rec.Result())
            if len(result['text']) > 0:
                if i % 5 == 0 and i % 10 != 0:
                    allText += result['text'] + '. '+"\n"
                    toSummary += result['text'] + '.'
                elif i % 10 == 0:
                    allText += result['text'] + '. ' + "\n" + "Timestamp: " + time.strftime("%H:%M:%S", time.gmtime(result['result'][len(result['result']) - 1]['end'])) + "\n"
                    toSummary += result['text'] + '.'
                else:
                    allText += result['text'] + '. '
                    toSummary += result['text'] + '. '
                percent = (result['result'][len(result['result']) - 1]['end']) / (len(total) / 1000) * 100
                # print(percent)
                percentLabel['text'] = "% Converted: " + "{:.2f}".format(percent)
                # print(result['result'][len(result['result'])-1]['end'])

    print("This is the final result")
    print(rec.FinalResult())
    print(allText)

    tool = language_tool_python.LanguageTool('en-US')
    matches = tool.check(allText)
    matches2 = tool.check(toSummary)
    print(matches)
    goodText = tool.correct(allText)
    goodSummary = tool.correct(toSummary)
    print(goodText)

    totalTimeLabel.destroy()
    percentLabel.destroy()

    global var
    print("Value of choice is: "+str(var.get()))
    if var.get() == 0 or var.get() == 1:
        completed(goodText)
    elif var.get() == 2:
        temp2 = Sum.Summarizer(goodSummary)
        summarizedText = temp2.summarize()
        completed(summarizedText)
    elif var.get() == 3:
        temp2 = Sum.Summarizer(goodSummary)
        summarizedText = temp2.summarize()
        completed(goodText+"\n"+summarizedText)


def start():
    threading.Thread(target=SpeechToText).start()


startButton = tk.Button(frame, text="Start conversion", command=start, background='#3b3a37', foreground='white', width='35', pady=5)


def popup_bonus():
    win = tk.Toplevel(top, height=200, width=200)
    win.wm_title("Save Confirmation")

    l = tk.Label(win, text="File saved successfully!")
    l.pack(side=tk.TOP)

    b = tk.Button(win, text="Okay", command=win.destroy)
    b.pack(side=tk.TOP)


def save_as(text):
    print(type(text))
    text_file = filedialog.asksaveasfilename(defaultextension=".pdf", initialdir="/", title="Save File", filetypes=(("PDF file", "*.pdf"), ("Word File", "*.docx"), ("Text file","*.txt")))

    if text_file:
        if text_file.find(".txt") != -1:
            text_file = open(text_file,'w')
            text_file.write(text)
            text_file.close()
        elif text_file.find(".pdf") != -1:
            my_doc = SimpleDocTemplate(text_file, pagesize=letter)
            flowables = []
            sample_style_sheet = getSampleStyleSheet()

            paragraphs = text.split('\n')
            paragraph_1 = Paragraph(os.path.basename(os.path.normpath(fileLabel['text']))+" to text", sample_style_sheet['Heading1'])
            flowables.append(paragraph_1)

            for paragraph in paragraphs:
                if paragraph.find("Timestamp") != -1:
                    new_paragraph = Paragraph(
                        paragraph,
                        sample_style_sheet['Italic']
                    )
                    flowables.append(new_paragraph)
                else:
                    new_paragraph = Paragraph(
                        paragraph,
                        sample_style_sheet['BodyText']
                    )
                    flowables.append(new_paragraph)
            my_doc.build(flowables)
        elif text_file.find("docx") != -1:
            document = Document()
            document.add_heading(os.path.basename(os.path.normpath(fileLabel['text']))+" to text", 0)
            paragraphs = text.split('\n')
            for paragraph in paragraphs:
                if paragraph.find("Timestamp") != -1:
                    p = document.add_paragraph()
                    p.add_run(paragraph).italic = True
                else:
                    document.add_paragraph(paragraph)
            document.save(text_file)
        popup_bonus()
        global filepath
        filepath = "None"
        fileLabel['text'] = "Selected file: " + filepath
        saveButton.pack_forget()
        completedLabel.pack_forget()
        fileButton.pack(side=tk.TOP)


completedLabel = tk.Label(frame, text="Conversion completed!", background='#242320', foreground='white', font=('Arial', 12), pady=5)
saveButton = tk.Button(frame, background='#3b3a37', foreground='white', width='35', pady=5)


def completed(text):
    completedLabel.pack(side=tk.TOP)
    saveButton.configure(text="Save file", command=lambda: save_as(text))
    saveButton.pack(side=tk.TOP)




madeByLabel = tk.Label(top, text="Made by Mustafa Muhammad and Adam Ding", foreground='white', background='#242320')
madeByLabel.place(relx=0.5, rely=0.98, anchor=tk.CENTER)
titleLabel = tk.Label(top, text="LectureAid",foreground='#3bb6e3', background='#242320', font=('Arial', 25))
titleLabel.place(relx=0.5,rely=0.05,anchor=tk.CENTER)
top.mainloop()

